"""
File Processing Utilities for AI Chatbot Hub

This module contains all file processing functions for handling various file types
including text documents, PDFs, Word documents, Excel files, CSV files, JSON files, and images.

Author: Pruthvirajsinh Zala
Version: 3.0
"""

import io
import json
import pandas as pd
from PIL import Image
import PyPDF2
import docx
from typing import Tuple, Union, Any


class FileProcessor:
    """Main class for handling different file types and processing them"""
    
    @staticmethod
    def get_file_details(uploaded_file) -> dict:
        """Extract basic file information"""
        return {
            "filename": uploaded_file.name,
            "filetype": uploaded_file.type,
            "filesize": uploaded_file.size
        }
    
    @staticmethod
    def process_text_file(uploaded_file) -> Tuple[str, dict]:
        """Process plain text files (.txt)"""
        try:
            content = str(uploaded_file.read(), "utf-8")
            return content, FileProcessor.get_file_details(uploaded_file)
        except Exception as e:
            return f"Error processing text file: {str(e)}", FileProcessor.get_file_details(uploaded_file)
    
    @staticmethod
    def process_pdf_file(uploaded_file) -> Tuple[str, dict]:
        """Process PDF files (.pdf)"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
            content = ""
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():  # Only add non-empty pages
                    content += f"--- Page {page_num} ---\n{page_text}\n\n"
            
            if not content.strip():
                content = "No readable text found in the PDF file."
            
            return content, FileProcessor.get_file_details(uploaded_file)
        except Exception as e:
            return f"Error processing PDF file: {str(e)}", FileProcessor.get_file_details(uploaded_file)
    
    @staticmethod
    def process_word_file(uploaded_file) -> Tuple[str, dict]:
        """Process Word documents (.docx)"""
        try:
            doc = docx.Document(io.BytesIO(uploaded_file.read()))
            content = ""
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    content += paragraph.text + "\n"
            
            # Process tables if any
            if doc.tables:
                content += "\n--- Tables in Document ---\n"
                for table_num, table in enumerate(doc.tables, 1):
                    content += f"\nTable {table_num}:\n"
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        content += row_text + "\n"
            
            if not content.strip():
                content = "No readable content found in the Word document."
            
            return content, FileProcessor.get_file_details(uploaded_file)
        except Exception as e:
            return f"Error processing Word file: {str(e)}", FileProcessor.get_file_details(uploaded_file)
    
    @staticmethod
    def process_excel_file(uploaded_file) -> Tuple[str, dict]:
        """Process Excel files (.xlsx, .xls)"""
        try:
            # Read all sheets
            excel_data = pd.read_excel(io.BytesIO(uploaded_file.read()), sheet_name=None)
            content = f"Excel file: {uploaded_file.name}\n"
            content += f"Number of sheets: {len(excel_data)}\n\n"
            
            for sheet_name, df in excel_data.items():
                content += f"--- Sheet: {sheet_name} ---\n"
                content += f"Shape: {df.shape[0]} rows, {df.shape[1]} columns\n\n"
                
                # Show first few rows
                content += "Data Preview:\n"
                content += df.head(10).to_string(index=True) + "\n\n"
                
                # Basic statistics for numeric columns
                numeric_cols = df.select_dtypes(include=['number']).columns
                if len(numeric_cols) > 0:
                    content += "Numeric Columns Summary:\n"
                    content += df[numeric_cols].describe().to_string() + "\n\n"
                
                # Column information
                content += "Column Information:\n"
                for col in df.columns:
                    content += f"- {col}: {df[col].dtype} (Non-null: {df[col].notna().sum()}/{len(df)})\n"
                content += "\n" + "="*50 + "\n\n"
            
            return content, FileProcessor.get_file_details(uploaded_file)
        except Exception as e:
            return f"Error processing Excel file: {str(e)}", FileProcessor.get_file_details(uploaded_file)
    
    @staticmethod
    def process_csv_file(uploaded_file) -> Tuple[str, dict]:
        """Process CSV files (.csv)"""
        try:
            df = pd.read_csv(io.BytesIO(uploaded_file.read()))
            content = f"CSV file: {uploaded_file.name}\n"
            content += f"Shape: {df.shape[0]} rows, {df.shape[1]} columns\n\n"
            
            # Data preview
            content += "Data Preview (first 10 rows):\n"
            content += df.head(10).to_string(index=True) + "\n\n"
            
            # Basic statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                content += "Numeric Columns Summary:\n"
                content += df[numeric_cols].describe().to_string() + "\n\n"
            
            # Column information
            content += "Column Information:\n"
            for col in df.columns:
                content += f"- {col}: {df[col].dtype} (Non-null: {df[col].notna().sum()}/{len(df)})\n"
            
            # Missing values summary
            missing_values = df.isnull().sum()
            if missing_values.sum() > 0:
                content += f"\nMissing Values:\n"
                for col, missing_count in missing_values[missing_values > 0].items():
                    content += f"- {col}: {missing_count} missing values\n"
            
            return content, FileProcessor.get_file_details(uploaded_file)
        except Exception as e:
            return f"Error processing CSV file: {str(e)}", FileProcessor.get_file_details(uploaded_file)
    
    @staticmethod
    def process_json_file(uploaded_file) -> Tuple[str, dict]:
        """Process JSON files (.json)"""
        try:
            json_data = json.loads(uploaded_file.read())
            content = f"JSON file: {uploaded_file.name}\n\n"
            
            # Pretty print JSON with proper formatting
            formatted_json = json.dumps(json_data, indent=2, ensure_ascii=False)
            content += "JSON Content:\n"
            content += formatted_json + "\n\n"
            
            # Basic structure analysis
            content += "Structure Analysis:\n"
            if isinstance(json_data, dict):
                content += f"- Root type: Dictionary with {len(json_data)} keys\n"
                content += f"- Keys: {list(json_data.keys())}\n"
            elif isinstance(json_data, list):
                content += f"- Root type: Array with {len(json_data)} items\n"
                if json_data and isinstance(json_data[0], dict):
                    content += f"- Sample keys in first item: {list(json_data[0].keys())}\n"
            else:
                content += f"- Root type: {type(json_data).__name__}\n"
            
            return content, FileProcessor.get_file_details(uploaded_file)
        except json.JSONDecodeError as e:
            return f"Invalid JSON format: {str(e)}", FileProcessor.get_file_details(uploaded_file)
        except Exception as e:
            return f"Error processing JSON file: {str(e)}", FileProcessor.get_file_details(uploaded_file)
    
    @staticmethod
    def process_image_file(uploaded_file) -> Tuple[Union[Image.Image, str], dict]:
        """Process image files (.jpg, .png, .gif, etc.)"""
        try:
            image = Image.open(io.BytesIO(uploaded_file.read()))
            
            # Create descriptive content about the image
            content = f"Image file: {uploaded_file.name}\n"
            content += f"Dimensions: {image.size[0]} x {image.size[1]} pixels\n"
            content += f"Format: {image.format}\n"
            content += f"Mode: {image.mode}\n"
            
            # Additional image info
            if hasattr(image, 'info') and image.info:
                content += f"Additional info: {image.info}\n"
            
            # For AI processing, return both the image object and description
            return image, FileProcessor.get_file_details(uploaded_file)
        except Exception as e:
            return f"Error processing image file: {str(e)}", FileProcessor.get_file_details(uploaded_file)


def process_uploaded_file(uploaded_file) -> Tuple[Any, dict]:
    """
    Main function to process different types of uploaded files
    
    Args:
        uploaded_file: Streamlit uploaded file object
    
    Returns:
        Tuple of (processed_content, file_details)
    """
    file_type = uploaded_file.type
    
    try:
        if file_type == "text/plain":
            return FileProcessor.process_text_file(uploaded_file)
        
        elif file_type == "application/pdf":
            return FileProcessor.process_pdf_file(uploaded_file)
        
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return FileProcessor.process_word_file(uploaded_file)
        
        elif file_type in ["application/vnd.ms-excel", 
                          "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
            return FileProcessor.process_excel_file(uploaded_file)
        
        elif file_type == "text/csv":
            return FileProcessor.process_csv_file(uploaded_file)
        
        elif file_type == "application/json":
            return FileProcessor.process_json_file(uploaded_file)
        
        elif file_type.startswith('image/'):
            return FileProcessor.process_image_file(uploaded_file)
        
        else:
            return (f"Unsupported file type: {file_type}", 
                   FileProcessor.get_file_details(uploaded_file))
            
    except Exception as e:
        return (f"Error processing file: {str(e)}", 
               FileProcessor.get_file_details(uploaded_file))


def upload_to_gemini(file_content, file_details) -> Tuple[Any, dict]:
    """
    Upload file content to Gemini for processing
    
    Args:
        file_content: Processed file content
        file_details: File metadata dictionary
    
    Returns:
        Tuple of (content, file_details)
    """
    try:
        if file_details["filetype"].startswith('image/'):
            # For images, return the image object as is
            return file_content, file_details
        else:
            # For text content, return as is
            return file_content, file_details
    except Exception as e:
        return f"Error uploading to Gemini: {str(e)}", file_details


def format_file_size(size_bytes: int) -> str:
    """
    Convert file size from bytes to human-readable format
    
    Args:
        size_bytes: File size in bytes
    
    Returns:
        Formatted file size string
    """
    if size_bytes == 0:
        return "0 B"
    
    size_names = ["B", "KB", "MB", "GB", "TB"]
    import math
    i = int(math.floor(math.log(size_bytes, 1024)))
    p = math.pow(1024, i)
    s = round(size_bytes / p, 2)
    return f"{s} {size_names[i]}"


def get_supported_file_types() -> dict:
    """
    Get dictionary of supported file types and their descriptions
    
    Returns:
        Dictionary mapping file types to descriptions
    """
    return {
        "text/plain": "Text files (.txt)",
        "application/pdf": "PDF documents (.pdf)",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "Word documents (.docx)",
        "application/vnd.ms-excel": "Excel files (.xls)",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "Excel files (.xlsx)",
        "text/csv": "CSV files (.csv)",
        "application/json": "JSON files (.json)",
        "image/jpeg": "JPEG images (.jpg, .jpeg)",
        "image/png": "PNG images (.png)",
        "image/gif": "GIF images (.gif)",
        "image/bmp": "BMP images (.bmp)",
        "image/webp": "WebP images (.webp)"
    }


def is_supported_file_type(file_type: str) -> bool:
    """
    Check if a file type is supported
    
    Args:
        file_type: MIME type of the file
    
    Returns:
        True if supported, False otherwise
    """
    supported_types = get_supported_file_types()
    return (file_type in supported_types or 
            file_type.startswith('image/'))
